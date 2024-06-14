from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE

doc = Document()
p = doc.add_paragraph()

url = "https://www.example.com"
text = "This is a clickable link"

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = paragraph.add_run(text)
    new_run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    new_run.font.underline = True

    hyperlink.append(new_run._r)
    paragraph._p.append(hyperlink)

    return hyperlink

add_hyperlink(p, text, url)

doc.save("demo_hyperlink.docx")
